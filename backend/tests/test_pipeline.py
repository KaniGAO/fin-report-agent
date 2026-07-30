"""端到端流水线轻量验证：样例解析 → 匹配 → 回填 → 校验输出"""
import sys
from pathlib import Path

BASE = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(BASE))

from docx import Document

from app.core.filler import analyze, fill_docx

SAMPLES = BASE / "samples"


def run():
    docx_path = SAMPLES / "template.docx"
    excel_paths = [SAMPLES / f"{n}.xlsx" for n in ("资产负债表", "利润表", "现金流量表")]

    reports, fill_plan, warnings = analyze(docx_path, excel_paths)
    assert not warnings, warnings
    assert len(reports) == 3, f"应识别 3 张表，实际 {len(reports)}"

    types = {r.statement_type for r in reports}
    assert types == {"资产负债表", "利润表", "现金流量表"}, types

    for r in reports:
        assert r.year_columns == ["2023", "2024"], (r.statement_type, r.year_columns)
        print(f"{r.statement_type}: matched={r.matched_count} unmatched={r.unmatched_count} "
              f"source={r.source_file}")
        # 每张表恰好 1 个模板独有项目应留空
        assert r.unmatched_count == 1, [i.docx_item for i in r.items if i.status == "unmatched"]
        for it in r.items:
            if it.status == "unmatched":
                print(f"  留空: {it.docx_item}")

    # 回填并校验关键数值
    out = SAMPLES / "filled.docx"
    fill_docx(docx_path, fill_plan, out)
    doc = Document(str(out))

    def cell_of(t_idx, item, col):
        table = doc.tables[t_idx]
        for row in table.rows:
            if row.cells[0].text.strip() == item:
                return row.cells[col].text.strip()
        raise AssertionError(f"未找到项目 {item}")

    # 资产负债表：货币资金 2023=12000 2024=15600（Excel 名「货 币 资 金」）
    assert cell_of(0, "货币资金", 1) == "12000", cell_of(0, "货币资金", 1)
    assert cell_of(0, "货币资金", 2) == "15600"
    # 同义词：所有者权益合计 ← 股东权益合计
    assert cell_of(0, "所有者权益合计", 2) == "35400"
    # 缺失项留空
    assert cell_of(0, "递延所得税资产", 1) == ""
    # 利润表：营业收入 ← 一、营业总收入
    assert cell_of(1, "营业收入", 1) == "45000"
    assert cell_of(1, "投资收益", 2) == ""
    # 现金流量表：负数保留
    assert cell_of(2, "投资活动产生的现金流量净额", 1) == "-3500"
    assert cell_of(2, "汇率变动对现金的影响", 1) == ""

    print(f"\ntest_pipeline: ALL PASS，输出 {out}")


def run_realistic():
    """用户真实格式端到端：年份表头在第 4 行、含报告期/报表类型/小节标题行"""
    docx_path = SAMPLES / "realistic_template.docx"
    excel_paths = [SAMPLES / "母公司资产负债表.xlsx"]

    reports, fill_plan, warnings = analyze(docx_path, excel_paths)
    assert len(reports) == 1, f"应识别 1 张表，实际 {len(reports)}"
    rep = reports[0]
    assert rep.statement_type == "资产负债表", rep.statement_type
    # 年份列顺序：最新年在前
    assert rep.year_columns == ["2026", "2025", "2024", "2023"], rep.year_columns

    names = {it.docx_item for it in rep.items}
    # 真实项目应被识别
    assert "货币资金" in names and "资产总计" in names, names
    # 元数据行与小节标题应被过滤
    assert "报告期" not in names, names
    assert "报表类型" not in names, names
    assert "流动资产：" not in names, names
    assert "非流动资产：" not in names, names
    # 全部匹配（模板与 Excel 同名）
    assert rep.unmatched_count == 0, [i.docx_item for i in rep.items if i.status != "matched"]

    out = SAMPLES / "realistic_filled.docx"
    fill_docx(docx_path, fill_plan, out)
    doc = Document(str(out))

    def cell_of(item, col):
        table = doc.tables[0]
        for row in table.rows:
            if row.cells[0].text.strip() == item:
                return row.cells[col].text.strip()
        raise AssertionError(f"未找到项目 {item}")

    def num(item, col):
        return float(cell_of(item, col))

    # 货币资金：2026=44.28, 2025=10.60, 2024=23.48, 2023=18.90（数值等价，允许 10.6）
    assert num("货币资金", 1) == 44.28, cell_of("货币资金", 1)
    assert num("货币资金", 2) == 10.60, cell_of("货币资金", 2)
    assert num("货币资金", 3) == 23.48, cell_of("货币资金", 3)
    assert num("货币资金", 4) == 18.90, cell_of("货币资金", 4)
    # 资产总计回填
    assert num("资产总计", 1) == 92.28, cell_of("资产总计", 1)

    print(f"\ntest_pipeline (realistic): ALL PASS，输出 {out}")


if __name__ == "__main__":
    run()
    run_realistic()

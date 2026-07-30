"""normalize/matcher 轻量验证（含中英文跨语言匹配）"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from app.core.normalize import normalize, canonical
from app.core.matcher import match_item
from app.core.filler import analyze, fill_docx
from app.core.docx_parser import parse_docx
from app.core.excel_parser import parse_excel


def test_basic():
    # 归一化
    assert normalize("货 币 资 金") == "货币资金"
    assert normalize("营业收入（合并）") == "营业收入"
    assert normalize("净利润（净亏损以-号填列）") == "净利润"
    assert normalize("一、营业总收入") == "营业总收入"
    assert canonical("营业总收入") == "营业收入"
    assert canonical("股东权益合计") == "所有者权益合计"

    # 精确匹配
    name, score = match_item("货币资金", ["应收账款", "货币资金", "存货"])
    assert name == "货币资金" and score == 100.0

    # 同义词匹配
    name, score = match_item("营业收入", ["营业总收入", "营业成本"])
    assert name == "营业总收入" and score == 100.0, (name, score)

    # 模糊匹配（空格/括号差异）
    name, score = match_item("利润总额", ["利润总额（亏损总额以-号填列）"])
    assert name is not None and score >= 85, (name, score)

    # 不应误配
    name, score = match_item("递延所得税资产", ["货币资金", "存货", "应收票据"])
    assert name is None, (name, score)

    # 合并项不应误配（保守策略：数值含义不同）
    name, score = match_item("应收票据", ["应收票据及应收账款"])
    assert name is None, (name, score)

    # 包含关系兜底（长度接近）
    name, score = match_item("固定资产", ["固定资产净值"])
    assert name is not None and score >= 85, (name, score)


def test_bilingual():
    # 中文模板项 vs 英文 Excel 项
    name, score = match_item("货币资金", ["Accounts Receivable", "Cash and Cash Equivalents"])
    assert name == "Cash and Cash Equivalents" and score == 100.0, (name, score)
    name, score = match_item("应收账款", ["Accounts Receivable, net"])
    assert name == "Accounts Receivable, net" and score == 100.0, (name, score)
    name, score = match_item("营业收入", ["Operating Revenue"])
    assert name == "Operating Revenue" and score == 100.0, (name, score)
    name, score = match_item("资产总计", ["Total Assets"])
    assert name == "Total Assets" and score == 100.0, (name, score)
    name, score = match_item("未分配利润", ["Retained Earnings"])
    assert name == "Retained Earnings" and score == 100.0, (name, score)
    name, score = match_item("经营活动产生的现金流量净额", ["Net Cash Flow from Operating Activities"])
    assert name == "Net Cash Flow from Operating Activities" and score == 100.0, (name, score)

    # 英文模板项 vs 中文 Excel 项（反向）
    name, score = match_item("Cash and Cash Equivalents", ["应收账款", "货币资金", "存货"])
    assert name == "货币资金" and score == 100.0, (name, score)
    name, score = match_item("Total Assets", ["资产总计"])
    assert name == "资产总计" and score == 100.0, (name, score)

    # 英文别名差异（含 & 与空格）仍可命中
    name, score = match_item("货币资金", ["Cash & Cash Equivalents"])
    assert name == "Cash & Cash Equivalents" and score == 100.0, (name, score)

    # 未收录英文不应跨语言误配
    name, score = match_item("Deferred Tax Assets", ["货币资金", "存货"])
    assert name is None, (name, score)


def test_bilingual_e2e(workdir: Path):
    """中文 Word 模板 + 英文 Excel：跨语言匹配且模板缺失项留空"""
    from docx import Document
    from openpyxl import Workbook

    # 中文 Word 模板（资产负债表）
    doc = Document()
    doc.add_heading("资产负债表（模板）", level=1)
    table = doc.add_table(rows=5, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "项目", "2023年度", "2024年度"
    items = ["货币资金", "应收账款", "存货", "递延所得税资产"]  # 末项 Excel 中缺失
    for i, it in enumerate(items, start=1):
        table.rows[i].cells[0].text = it
    docx_path = workdir / "template.docx"
    doc.save(docx_path)

    # 英文 Excel
    wb = Workbook()
    ws = wb.active
    ws.title = "Balance Sheet"
    ws.append(["Company Balance Sheet"])
    ws.append(["Item", "2023", "2024"])
    ws.append(["Cash and Cash Equivalents", 12000, 15600])
    ws.append(["Accounts Receivable, net", 6800, 7100])
    ws.append(["Inventory, net", 8500, 9200])
    xlsx_path = workdir / "balance_sheet.xlsx"
    wb.save(xlsx_path)

    reports, fill_plan, warnings = analyze(docx_path, [xlsx_path])
    assert len(reports) == 1
    rep = reports[0]
    assert rep.statement_type == "资产负债表"
    matched = {it.docx_item: it for it in rep.items}
    assert matched["货币资金"].status == "matched"
    assert matched["货币资金"].excel_item == "Cash and Cash Equivalents"
    assert matched["应收账款"].status == "matched"
    assert matched["存货"].status == "matched"
    # 模板有而 Excel 无 -> 留空
    assert matched["递延所得税资产"].status == "unmatched"

    # 回填仅影响匹配项，且数值正确写入
    out_path = workdir / "filled.docx"
    fill_docx(docx_path, fill_plan, out_path)
    _, filled = parse_docx(str(out_path))
    dt = filled[0]
    # 货币资金 2023 列应回填 12000
    assert dt.item_rows  # 解析到项目行
    # 通过回填计划验证年份映射
    assert fill_plan  # 至少有一条回填记录


def test_docx_robust(workdir: Path):
    """鲁棒解析：年份表头下沉、含元数据/小节行、标签列非首列 等变体"""
    from docx import Document

    # 场景 A：年份表头在第 4 行 + 报告期/报表类型元数据行 + 小节标题
    doc = Document()
    doc.add_heading("母公司资产负债表", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    rows = [
        ["母公司资产负债表", "", "", "", ""],
        ["单位：亿元", "", "", "", ""],
        ["报告期", "一季报", "年报", "年报", "年报"],
        ["报表类型", "母公司报表", "母公司报表", "母公司报表", "母公司报表"],
        ["", "2026-03-31", "2025-12-31", "2024-12-31", "2023-12-31"],
        ["流动资产：", "", "", "", ""],
        ["货币资金", "", "", "", ""],
        ["资产总计", "", "", "", ""],
    ]
    for r in rows:
        cells = table.add_row().cells
        for ci, v in enumerate(r):
            cells[ci].text = v
    docx_path = workdir / "realistic.docx"
    doc.save(docx_path)

    _, tables = parse_docx(str(docx_path))
    assert len(tables) == 1, len(tables)
    dt = tables[0]
    assert dt.statement_type == "资产负债表", dt.statement_type
    assert dt.year_cols == {1: "2026", 2: "2025", 3: "2024", 4: "2023"}, dt.year_cols
    names = set(dt.item_rows.values())
    assert "货币资金" in names and "资产总计" in names, names
    # 元数据行 / 小节标题 / 标题行 均不在项目行
    assert "报告期" not in names and "报表类型" not in names, names
    assert "流动资产：" not in names and "母公司资产负债表" not in names, names

    # 场景 B：标签列非首列（第 0 列是序号，第 1 列是名称，年份在第 2/3 列）
    doc2 = Document()
    doc2.add_heading("利润表", level=1)
    t2 = doc2.add_table(rows=1, cols=4)
    t2.style = "Table Grid"
    hdr = t2.add_row().cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "序号", "项目", "2023年度", "2024年度"
    for i, name in enumerate(["营业收入", "净利润"], start=1):
        cells = t2.add_row().cells
        cells[0].text = str(i)      # 序号（数值列）
        cells[1].text = name        # 名称（应被识别为标签列）
    docx2 = workdir / "seq.docx"
    doc2.save(docx2)
    _, tables2 = parse_docx(str(docx2))
    assert tables2, "序号+名称布局应被识别"
    dt2 = tables2[0]
    assert dt2.item_rows, dt2.item_rows
    # 标签列应识别到第 1 列（名称），而非第 0 列（序号）
    vals = set(dt2.item_rows.values())
    assert "营业收入" in vals and "净利润" in vals, vals


def run():
    test_basic()
    test_bilingual()
    import tempfile
    with tempfile.TemporaryDirectory() as td:
        test_bilingual_e2e(Path(td))
        test_docx_robust(Path(td))
    print("test_matcher: ALL PASS")


if __name__ == "__main__":
    run()

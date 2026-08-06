"""单位换算回归测试：复现「模板标亿元、Excel 是万元」错位问题"""
import sys
import tempfile
from pathlib import Path

BASE = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(BASE))

from docx import Document
from openpyxl import Workbook

from app.core.filler import analyze, fill_docx
from app.core.units import normalize_unit, convert_value


def test_convert_value_basic():
    # 万元 -> 亿元 除以 10000（带千分位、两位小数）
    assert convert_value("17089900", "万元", "亿元") == "1,708.99"
    # 亿元 -> 万元 乘以 10000
    assert convert_value("1708.99", "亿元", "万元") == "17,089,900"
    # 同单位原样返回
    assert convert_value("1234", "亿元", "亿元") == "1234"
    # 单位未知 -> 保守原样返回
    assert convert_value("1234", None, "亿元") == "1234"
    assert convert_value("1234", "万元", None) == "1234"
    # 负数保留
    assert convert_value("-3500", "万元", "亿元") == "-0.35"
    # 千分位输入
    assert convert_value("1,708,9900", "万元", "亿元") == "1,708.99"
    print("test_convert_value_basic: ALL PASS")


def test_normalize_unit():
    assert normalize_unit("贵州茅台(600519) 利润表（单位：亿元）") == "亿元"
    assert normalize_unit("资产负债表（单位：万元）") == "万元"
    assert normalize_unit("报表（千元）") == "千元"
    assert normalize_unit("金额单位：百万元") == "百万元"
    assert normalize_unit("营业收入", "2023", "2024") is None
    print("test_normalize_unit: ALL PASS")


def test_excel_wan_to_doc_yi():
    """完整链路：Excel 数据以万元为单位，Word 模板标注亿元，应自动换算"""
    tmp = Path(tempfile.mkdtemp())

    # 1) Excel：标题含「单位：万元」，营业收入 2023=17089900 万元
    wb = Workbook()
    ws = wb.active
    ws.title = "利润表"
    ws.append(["利润表（单位：万元）"])
    ws.append(["项目", "2023", "2024"])
    ws.append(["营业收入", 17089900, 19800000])
    ws.append(["净利润", 8000000, 9200000])
    xlsx = tmp / "利润表.xlsx"
    wb.save(xlsx)

    # 2) Word：表格前段落标注「单位：亿元」
    doc = Document()
    doc.add_paragraph("XX公司 利润表（单位：亿元）")
    t = doc.add_table(rows=3, cols=3)
    t.rows[0].cells[0].text = "项目"
    t.rows[0].cells[1].text = "2023"
    t.rows[0].cells[2].text = "2024"
    t.rows[1].cells[0].text = "营业收入"
    t.rows[2].cells[0].text = "净利润"
    docx = tmp / "template.docx"
    doc.save(docx)

    reports, fill_plan, warnings = analyze(docx, [xlsx])
    assert len(reports) == 1
    # 应触发换算警告
    assert any("亿元" in w and "万元" in w for w in warnings), warnings

    out = tmp / "filled.docx"
    fill_docx(docx, fill_plan, out)
    d = Document(str(out))
    table = d.tables[0]

    def cell_of(item, col):
        for row in table.rows:
            if row.cells[0].text.strip() == item:
                return row.cells[col].text.strip()
        raise AssertionError(f"未找到 {item}")

    # 17089900 万元 -> 1708.99 亿元
    assert cell_of("营业收入", 1) == "1,708.99", cell_of("营业收入", 1)
    assert cell_of("营业收入", 2) == "1,980", cell_of("营业收入", 2)
    assert cell_of("净利润", 1) == "800", cell_of("净利润", 1)
    print("test_excel_wan_to_doc_yi: ALL PASS")


if __name__ == "__main__":
    test_convert_value_basic()
    test_normalize_unit()
    test_excel_wan_to_doc_yi()
    print("\ntest_units: ALL PASS")

"""Word 模板表格解析：自适应识别报表类型、年份列、项目行

相比旧版，本解析器不再假设「年份表头必在前 3 行、项目名必在第 0 列」，
而是：
  - 扫描整张表寻找含年份最多的行作为表头（支持单表堆叠多张报表）；
  - 自动探测项目名列（兼容「序号列 + 名称列」等布局）；
  - 过滤「报告期/报表类型/单位」等元数据行与「流动资产：」等小节标题。
"""
from __future__ import annotations

from dataclasses import dataclass, field

from docx import Document
from docx.table import Table

from .excel_parser import (
    detect_statement_type,
    extract_year,
    _looks_numeric,
    _is_metadata_label,
)


@dataclass
class DocxTable:
    table_index: int                     # doc.tables 中的序号
    statement_type: str
    header_row: int                      # 表头行号（表内绝对行号）
    year_cols: dict[int, str] = field(default_factory=dict)   # {列号: 年份}
    item_rows: dict[int, str] = field(default_factory=dict)   # {绝对行号: 项目名}


def _table_texts(table: Table, max_rows: int = 3) -> list[str]:
    texts = []
    for row in table.rows[:max_rows]:
        for cell in row.cells:
            texts.append(cell.text)
    return texts


def _paragraphs_before(doc: Document, table: Table, n: int = 5) -> list[str]:
    """取表格前面最近的 n 个非空段落文本（用于识别报表类型，如标题「资产负债表」）"""
    texts = []
    for block in doc.element.body:
        if block is table._tbl:
            break
        if block.tag.endswith("}p"):
            t = "".join(node.text or "" for node in block.iter() if node.tag.endswith("}t"))
            if t.strip():
                texts.append(t.strip())
    return texts[-n:]


def _is_section_header(name: str) -> bool:
    """小节标题：以 ：/: / 、 / ； / ; / ， 结尾的辅助行，不应当作项目"""
    name = name.strip()
    if not name:
        return False
    return name[-1] in "：:；;、"


def _year_cells(cells: list[str]) -> dict[int, str]:
    """从一行单元格中，统计列号 >=1 内含年份的单元格 -> {列号: 年份}"""
    cols: dict[int, str] = {}
    for ci in range(1, len(cells)):
        y = extract_year(cells[ci])
        if y:
            cols[ci] = y
    return cols


def _detect_label_col(
    all_rows: list[list[str]], year_cols: dict[int, str], start: int, end: int, n_cols: int
) -> int | None:
    """选取项目名列：优先年份列左侧、其次右侧，取「含文本（非数值）最多」的列。
    兼容「序号列 + 名称列」布局，且不会把纯数值的序号列误当名称列。"""
    if not year_cols:
        return None
    min_col = min(year_cols)
    max_col = max(year_cols)
    left = list(range(0, min_col)) if min_col >= 1 else []
    right = list(range(max_col + 1, n_cols))
    best_col, best_score = None, 0
    for ci in left + right:
        score = 0
        for ri in range(start, end):
            if ri >= len(all_rows):
                break
            cells = all_rows[ri]
            if ci >= len(cells):
                continue
            t = cells[ci].strip()
            if t and not _looks_numeric(t):
                score += 1
        if score > best_score:
            best_score, best_col = score, ci
    return best_col if best_score > 0 else None


def parse_docx(path: str) -> tuple[Document, list[DocxTable]]:
    """解析 docx，返回 (Document 对象, 可填写的表格列表)。
    Document 对象保留引用以便后续 filler 直接回填并保存。"""
    doc = Document(path)
    results: list[DocxTable] = []

    for ti, table in enumerate(doc.tables):
        if not table.rows:
            continue
        n_rows = len(table.rows)
        all_rows = [[c.text for c in row.cells] for row in table.rows]
        n_cols = max((len(r) for r in all_rows), default=0)

        # 1) 候选年份表头：整表扫描，含年份(列>=1)单元格最多的行。
        #    要求 >=2 个年份以提高鲁棒性；若无则放宽到 >=1 个。
        header_candidates: list[tuple[int, dict[int, str]]] = []
        for ri in range(n_rows):
            yc = _year_cells(all_rows[ri])
            if len(yc) >= 2:
                header_candidates.append((ri, yc))
        if not header_candidates:
            for ri in range(n_rows):
                yc = _year_cells(all_rows[ri])
                if yc:
                    header_candidates.append((ri, yc))
        if not header_candidates:
            continue  # 该表无年份表头，跳过

        # 报表类型：从最近的段落倒序判断，仍未知时用表头行文本兜底
        st_type = "未知"
        for text in reversed(_paragraphs_before(doc, table)):
            st_type = detect_statement_type(text)
            if st_type != "未知":
                break
        if st_type == "未知":
            st_type = detect_statement_type(*_table_texts(table, max_rows=1))

        # 2) 逐个候选表头切分区间，支持「单表堆叠多张报表」
        for idx, (h, year_cols) in enumerate(header_candidates):
            next_h = header_candidates[idx + 1][0] if idx + 1 < len(header_candidates) else n_rows
            label_col = _detect_label_col(all_rows, year_cols, h + 1, next_h, n_cols)
            if label_col is None:
                continue

            item_rows: dict[int, str] = {}
            for ri in range(h + 1, next_h):
                cells = all_rows[ri]
                if label_col >= len(cells):
                    continue
                name = cells[label_col].strip()
                if not name:
                    continue
                # 过滤元数据行与小节标题
                if _is_section_header(name) or _is_metadata_label(name):
                    continue
                item_rows[ri] = name

            if item_rows:
                results.append(
                    DocxTable(
                        table_index=ti,
                        statement_type=st_type,
                        header_row=h,
                        year_cols=year_cols,
                        item_rows=item_rows,
                    )
                )

    return doc, results

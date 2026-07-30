"""生成演示样例：1 个 Word 模板 + 3 个 Excel 财务报表
特点：项目顺序打乱、命名有差异、部分模板项目 Excel 中缺失
"""
import sys
from pathlib import Path

from docx import Document
from openpyxl import Workbook

OUT = Path(__file__).resolve().parent.parent / "samples"
OUT.mkdir(exist_ok=True)

YEARS = ["2023", "2024"]

# ---------- Word 模板定义：项目固定，顺序与 Excel 不同 ----------
DOCX_TABLES = {
    "资产负债表": [
        "货币资金", "应收账款", "存货", "流动资产合计",
        "固定资产", "无形资产", "资产总计",
        "短期借款", "应付账款", "流动负债合计",
        "长期借款", "负债合计",
        "实收资本", "未分配利润", "所有者权益合计",
        "负债和所有者权益总计",
        "递延所得税资产",  # Excel 中缺失，应留空
    ],
    "利润表": [
        "营业收入", "营业成本", "税金及附加",
        "销售费用", "管理费用", "研发费用", "财务费用",
        "营业利润", "利润总额", "所得税费用", "净利润",
        "投资收益",  # Excel 中缺失，应留空
    ],
    "现金流量表": [
        "经营活动产生的现金流量净额",
        "投资活动产生的现金流量净额",
        "筹资活动产生的现金流量净额",
        "现金及现金等价物净增加额",
        "期末现金及现金等价物余额",
        "汇率变动对现金的影响",  # Excel 中缺失，应留空
    ],
}

# ---------- Excel 数据：命名有差异、顺序打乱 ----------
EXCEL_DATA = {
    "资产负债表": {
        "sheet": "资产负债表",
        "rows": [
            ("存货净额", 8500, 9200),
            ("货 币 资 金", 12000, 15600),          # 含空格
            ("应收账款净额", 6800, 7100),
            ("流动资产合计", 27300, 31900),
            ("固定资产净值", 18000, 17200),
            ("无形资产净额", 3200, 3000),
            ("资产总计", 48500, 52100),
            ("应付账款", 5600, 6200),
            ("短期借款", 4000, 3500),
            ("流动负债合计", 9600, 9700),
            ("长期借款", 8000, 7000),
            ("负债合计", 17600, 16700),
            ("实收资本（或股本）", 10000, 10000),      # 括号差异
            ("未分配利润", 12900, 17400),
            ("股东权益合计", 30900, 35400),           # 同义词
            ("负债及所有者权益总计", 48500, 52100),    # 同义词
        ],
    },
    "利润表": {
        "sheet": "利润表",
        "rows": [
            ("一、营业总收入", 45000, 52000),          # 序号 + 同义词
            ("营业总成本", 30000, 33500),
            ("税金及附加", 450, 520),
            ("销售费用", 3200, 3600),
            ("管理费用", 2800, 3000),
            ("研发费用", 1500, 1900),
            ("财务费用", 600, 480),
            ("营业利润", 6450, 9000),
            ("利润总额（亏损总额以-号填列）", 6400, 8950),
            ("所得税", 960, 1340),                    # 同义词
            ("净利润（净亏损以-号填列）", 5440, 7610),
        ],
    },
    "现金流量表": {
        "sheet": "现金流量表",
        "rows": [
            ("经营活动现金流量净额", 7200, 8900),       # 同义词
            ("投资活动现金流量净额", -3500, -2800),
            ("筹资活动现金流量净额", -1200, -2400),
            ("现金及现金等价物净增加（减少）额", 2500, 3700),
            ("期末现金及现金等价物余额", 12000, 15700),
        ],
    },
}


def gen_docx():
    doc = Document()
    doc.add_heading("XX 公司财务分析报告（模板）", level=0)
    doc.add_paragraph("以下三张报表由财务部按年度数据填写，单位：万元。")

    for st_name, items in DOCX_TABLES.items():
        doc.add_heading(f"附表：{st_name}", level=1)
        table = doc.add_table(rows=1 + len(items), cols=1 + len(YEARS))
        table.style = "Table Grid"
        # 表头
        hdr = table.rows[0].cells
        hdr[0].text = "项目"
        for i, y in enumerate(YEARS, start=1):
            hdr[i].text = f"{y}年度"
        # 项目行（数值列留空）
        for ri, item in enumerate(items, start=1):
            table.rows[ri].cells[0].text = item
        doc.add_paragraph("")

    path = OUT / "template.docx"
    doc.save(path)
    print(f"生成 {path}")


def gen_excels():
    for st_name, cfg in EXCEL_DATA.items():
        wb = Workbook()
        ws = wb.active
        ws.title = cfg["sheet"]
        ws.append([f"XX公司{st_name}"])
        ws.append(["项目", f"{YEARS[0]}年12月31日", f"{YEARS[1]}年12月31日"])
        for row in cfg["rows"]:
            ws.append(list(row))
        path = OUT / f"{st_name}.xlsx"
        wb.save(path)
        print(f"生成 {path}")


# ---------- 「用户格式」样例：年份表头在第 4 行、含报告期/报表类型元数据行、含小节标题 ----------
# 模拟用户真实「母公司资产负债表」：单位行/报告期/报表类型 占位，年份行下沉，项目名在第 0 列
REAL_YEARS = ["2026-03-31", "2025-12-31", "2024-12-31", "2023-12-31"]
REAL_ITEMS = [
    ("流动资产：", None),                       # 小节标题（应被过滤）
    ("货币资金", ("货币资金", 44.28, 10.60, 23.48, 18.90)),
    ("应收账款", ("应收账款", 12.30, 9.80, 8.10, 7.50)),
    ("存货", ("存货", 5.60, 4.20, 3.90, 3.10)),
    ("流动资产合计", ("流动资产合计", 62.18, 24.60, 35.48, 29.50)),
    ("非流动资产：", None),                     # 小节标题
    ("固定资产", ("固定资产", 30.10, 28.40, 27.00, 25.50)),
    ("资产总计", ("资产总计", 92.28, 52.00, 62.48, 55.00)),
]


def gen_realistic():
    """生成用户真实格式 Word 模板 + 对应中文 Excel，用于验证鲁棒解析器。"""
    # Word 模板（数值列留空，待回填）
    doc = Document()
    doc.add_heading("母公司资产负债表", level=1)   # 用于报表类型识别
    table = doc.add_table(rows=1, cols=1 + len(REAL_YEARS))
    table.style = "Table Grid"
    rows = [
        ["母公司资产负债表", "", "", "", ""],       # 标题行（无年份）
        ["单位：亿元", "", "", "", ""],             # 单位行（无年份）
        ["报告期", "一季报", "年报", "年报", "年报"],      # 元数据行
        ["报表类型", "母公司报表", "母公司报表", "母公司报表", "母公司报表"],  # 元数据行
        ["", *REAL_YEARS],                          # 真正的年份表头（第 4 行）
    ]
    for r in rows:
        cells = table.add_row().cells
        for ci, v in enumerate(r):
            cells[ci].text = v

    for label, _ in REAL_ITEMS:
        cells = table.add_row().cells
        cells[0].text = label
        for ci in range(1, len(REAL_YEARS) + 1):
            cells[ci].text = ""   # 数值列留空

    path = OUT / "realistic_template.docx"
    doc.save(path)
    print(f"生成 {path}")

    # 对应 Excel（项目名与模板一致，年份格式一致）
    wb = Workbook()
    ws = wb.active
    ws.title = "母公司资产负债表"
    ws.append(["母公司资产负债表"])
    ws.append(["项目", *REAL_YEARS])
    for _, payload in REAL_ITEMS:
        if payload is not None:
            ws.append(list(payload))
    xlsx = OUT / "母公司资产负债表.xlsx"
    wb.save(xlsx)
    print(f"生成 {xlsx}")


if __name__ == "__main__":
    gen_docx()
    gen_excels()
    gen_realistic()
